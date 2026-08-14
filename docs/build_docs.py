#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""生成《SeaUI 软件框架说明》DOCX 与 PDF。

排版遵循 documents 技能的 standard_business_brief 预设：
  US Letter / 1in 边距 / Calibri+微软雅黑 11pt / 行距 1.10 / 段后 6pt
  H1 16pt #2E74B5 / H2 13pt #2E74B5 / H3 12pt #1F4D78
  列表：标记 0.25in、文字 0.5in、悬挂 0.25in
  表格：9360 DXA 固定宽度、缩进 120 DXA、单元格边距 80/80/120/120、
        表头填充 #F2F4F7
"""

from __future__ import annotations

from pathlib import Path
from typing import Any

DOCX_PATH = Path(__file__).resolve().parent / "SeaUI_软件框架说明.docx"
PDF_PATH = Path(__file__).resolve().parent / "SeaUI_软件框架说明.pdf"

HEADING_BLUE = "#2E74B5"
HEADING_DARK_BLUE = "#1F4D78"
INK_BLUE = "#0B2545"
HEADER_FILL = "#F2F4F7"
GRAY = "#595959"
BORDER_GRAY = "#C9CDD3"

# ============================================================================
# 共享内容
# ============================================================================
CONTENT: list[tuple[str, Any]] = [
    ("title", "SeaUI 海参检测及吸捕机器人控制系统"),
    ("subtitle", "软件框架说明 · 软件（PC）→ RDK X5 → Pixhawk 2.4.8"),
    ("meta", "版本 v2.2.0    编制日期 2026-08-15    密级 内部技术文档"),

    ("h1", "1. 总体架构"),
    ("para", "系统由三层组成：PC 端的 Flutter 桌面软件负责交互与展示；PC 端 Python 后端负责桥接、数据库与鉴权；地瓜机器人 RDK X5（Sunrise 5）承担摄像头采集、BPU 上 YOLO11 海参分割推理、传感器读取与 Pixhawk 控制。三层之间通过网线直连的一条 WebSocket 通道交换视频、遥测与控制指令。"),
    ("code", (
        "PC（Windows）                                                     \n"
        "  Flutter 桌面界面 <--ws://127.0.0.1:8765--> Python 后端（桥接）      \n"
        "  http://127.0.0.1:5000（登录 / 管理员 / 传感器 / 日志 REST）         \n"
        "                              |  SQLite 数据库（用户·会话·传感器·日志）\n"
        "                              |                                       \n"
        "=========================== 网线直连 ===============================\n"
        "                              | ws://192.168.127.10:8080            \n"
        "RDK X5（Ubuntu 22.04）                                               \n"
        "  MIPI/USB 摄像头 -> BPU YOLO11 分割 -> 标注 JPEG 视频流               \n"
        "  VEML7700 / MS5837 / DS18B20 / LO81MTW -> 传感器遥测                  \n"
        "  WebSocket 服务器（视频 + 遥测 + 命令）                               \n"
        "                              | MAVLink /dev/ttyACM0                 \n"
        "Pixhawk 2.4.8（ArduSub / PX4）                                        \n"
        "  MAIN1-8 推进器 · AUX1-2 吸捕电机 · AUX3 舵机                         "
    )),

    ("h1", "2. PC 端软件（Flutter + Python 桥接）"),
    ("table", {
        "headers": ["组件", "位置", "职责"],
        "widths": [1800, 2600, 4960],
        "rows": [
            ["Flutter 桌面界面", "rov_flutter/", "登录、主控、控制操作、数据分析、设置"],
            ["本地后端", "backend/app.py", "桥接、鉴权、REST、命令转发、死区看门狗"],
            ["RDK 客户端", "backend/rdk_client.py", "连接 RDK X5 WebSocket，断线自动重连"],
            ["数据库", "backend/database.py", "SQLite：用户、会话、传感器、控制日志、设置"],
            ["REST API", "http://127.0.0.1:5000", "登录、管理员、传感器、日志、命令"],
        ],
    }),
    ("para", "后端支持三种运行模式（环境变量 ROV_BACKEND_MODE）："),
    ("bullets", [
        "rdk（默认）：视频与 AI 全部来自 RDK X5，PC 不再运行 ONNX。",
        "local：旧版兼容模式，使用 PC 摄像头与 best.onnx 推理，仅用于无板卡开发调试。",
        "sim：合成视频帧与遥测，无任何硬件也能完成界面联调。",
    ]),

    ("h1", "3. RDK X5 端（rdkx5/，代码已逐文件标注 [RDK X5 side]）"),
    ("table", {
        "headers": ["文件", "职责"],
        "widths": [2200, 7160],
        "rows": [
            ["gateway.py", "主程序，装配所有组件并处理退出清理"],
            ["stream_server.py", "WebSocket 视频 / 遥测 / 命令通道"],
            ["vision.py", "srcampy MIPI 采集、hobot_dnn BPU 分割、OpenCV 标注、JPEG 编码"],
            ["pixhawk_link.py", "MAVLink 控制、状态遥测、死区看门狗"],
            ["sensors.py", "VEML7700 / MS5837 / DS18B20 / LO81MTW 驱动"],
            ["yolo11_seg_rdk.py", "量化 BIN 模型推理脚本（来自模型转换仓库）"],
            ["YOLO11_LBL.bin", "海参 YOLO11 分割量化模型"],
            ["config.yaml", "网络、摄像头、模型、传感器、Pixhawk 接线参数"],
            ["PROTOCOL.md", "PC 与 RDK X5 的消息协议"],
            ["check_hardware.py", "实机一键自检：网络/设备/依赖/传感器/Pixhawk/BPU/端口"],
        ],
    }),

    ("h1", "4. 通信协议（一条 WebSocket，端口 8080）"),
    ("h2", "4.1 下行消息（RDK X5 → PC）"),
    ("table", {
        "headers": ["类型", "关键字段", "说明"],
        "widths": [1500, 4200, 3660],
        "rows": [
            ["hello", "device, version, caps", "连接建立后发送一次"],
            ["frame", "seq, width, height, jpeg(base64), detections[], inference_ms, fps", "标注后视频帧，约 15 fps"],
            ["telemetry", "sensors{}, pixhawk{...}, link{fps}", "传感器与飞控状态，约 5 Hz"],
            ["log / ack", "level / command, success", "运行日志与命令确认"],
        ],
    }),
    ("h2", "4.2 上行命令（PC → RDK X5）"),
    ("table", {
        "headers": ["命令", "参数", "说明"],
        "widths": [1800, 3000, 4560],
        "rows": [
            ["move", "axes{surge,sway,heave,roll,pitch,yaw}, deadman_ms", "六自由度速度，范围 -1..1"],
            ["stop / arm / disarm / set_mode", "- / force / mode", "回中、解锁、锁定、切模式"],
            ["suction / servo", "power_percent / channel, pwm", "吸捕电机 0-100%，舵机直通 PWM"],
            ["light_on/off / sonar_on/off / laser_on/off", "-", "辅助设备开关"],
            ["emergency_stop / snapshot / reset_position", "-", "急停、抓拍、坐标归零"],
        ],
    }),
    ("note", "安全约定：move 必须携带 deadman_ms（默认 1000 ms），PC 每 100 ms 重发一次；RDK X5 超时自动输出中立值。启动时 Pixhawk 保持未解锁，只有收到 arm 才解锁。"),

    ("h1", "5. Pixhawk 控制链"),
    ("para", "控制路径为 SeaUI 界面 → PC 后端 → RDK X5 gateway → MAVLink → Pixhawk 2.4.8。pixhawk_link.py 提供两种控制方式，由 config.yaml 的 pixhawk.control_mode 选择："),
    ("bullets", [
        "manual_control（推荐）：ArduSub MANUAL_CONTROL 虚拟摇杆（x 前后、y 左右、z 垂直、r 偏航），保留飞控稳定回路与混控；吸捕电机与舵机仍通过 DO_SET_SERVO 输出 AUX。",
        "servo_pwm：MAV_CMD_DO_SET_SERVO 直通 PWM，兼容 PX4 与任意固件，推进器混控由上位机完成。",
    ]),
    ("para", "原 GitHub 仓库 Sunrise5-Based-Sea-Cucumber-Inspection-and-Suction-Harvest-Robot 的 pixhawk_mavlink.py 已同步优化，新增 arm、set_mode、manual_control、drain_messages、telemetry_snapshot 等接口（本地 reference/ 克隆中）。"),

    ("h1", "6. 传感器回传（全部接在 RDK X5 上）"),
    ("table", {
        "headers": ["传感器", "接口", "输出字段"],
        "widths": [2200, 1800, 5360],
        "rows": [
            ["VEML7700 光照 ×2", "I2C", "lux, als_raw, white_raw"],
            ["MS5837-30BA 压力/深度", "I2C", "pressure_mbar, temperature_c, depth_m"],
            ["DS18B20 水温 ×2", "1-Wire sysfs", "temperature_c"],
            ["LO81MTW 超声波 ×2", "USB 串口", "distance_m"],
        ],
    }),
    ("para", "读数进入 telemetry.sensors → PC 后端写入 SQLite sensor_readings 表 → Flutter 主控页与操作页实时显示；数据分析页从 GET /api/sensors 读取历史并按分钟聚合，绘制深度、水温、光照、压强曲线。无板卡时 sim 模式会生成合成遥测，用于演示完整链路。"),

    ("h1", "7. 数据库与登录拦截"),
    ("bullets", [
        "SQLite 表：users（管理员账号）、sessions（登录会话）、sensor_readings（传感器数据）、control_logs（控制日志）、settings（运行设置）。",
        "密码使用 PBKDF2-HMAC-SHA256 加盐哈希，数据库中不保存明文。",
        "超级管理员在首次启动时自动创建：用户名 zmm，密码 Zmm771023。",
        "登录走 POST /api/login；用户名或密码缺失、错误一律返回 401，Flutter 登录页直接拦截，不再出现“访客放行”路径。",
        "用户管理接口需要 Bearer token，且角色必须为 super_admin 或 admin；控制命令全部写入 control_logs。",
        "Flutter 管理员面板的用户列表/增删、操作日志均以数据库为权威来源。",
    ]),

    ("h1", "8. 部署与运行"),
    ("h2", "8.1 PC 端"),
    ("code", "双击 open_seaUI.bat（自动拉起后端并打开桌面程序）\n需要重新构建时：open_seaUI.bat /rebuild"),
    ("h2", "8.2 RDK X5 端"),
    ("code", (
        "# PC 上传输代码\n"
        "scp -r rdkx5 sunrise@192.168.127.10:/home/sunrise/seaUI_rdk\n"
        "# 板卡上运行（默认用户名/密码 sunrise）\n"
        "ssh sunrise@192.168.127.10\n"
        "cd /home/sunrise/seaUI_rdk && ./run_robot.sh"
    )),
    ("note", "板卡网口默认静态 IP 192.168.127.10；PC 网口需配置为同网段（如 192.168.127.100 / 255.255.255.0）。"),
    ("h2", "8.3 无硬件联调"),
    ("code", "set ROV_BACKEND_MODE=sim\npython backend\\app.py"),

    ("h1", "9. 验证情况"),
    ("bullets", [
        "backend/tests 共 11 项测试全部通过：超级管理员初始化与鉴权拦截、会话、用户管理、传感器/控制日志入库、REST、UI WebSocket、假 RDK 网关回环（视频 + 遥测转发）。",
        "一键脚本实测：桌面程序启动、Python 后端自动拉起、8765/5000 端口监听、错误密码返回 401、超级管理员登录成功。",
        "RDK X5 端代码全部通过 py_compile 语法检查；摄像头、BPU、传感器与 Pixhawk 控制需在实机板卡上按 rdkx5/README.md 逐项验证。",
        "YOLO11_LBL.bin 的 SHA256 与上游仓库记录一致；rdkx5/check_hardware.py 提供实机一键自检。",
        "sim 模式实测：合成传感器遥测落库并经 GET /api/sensors 返回，供数据分析页消费。",
    ]),

    ("h1", "10. 参考资料"),
    ("bullets", [
        "D-Robotics RDK 资料中心：https://developer.d-robotics.cc/rdk_doc_center/",
        "Sunrise5 海参检测与吸捕机器人仓库：https://github.com/bertholetnadine159-prog/Sunrise5-Based-Sea-Cucumber-Inspection-and-Suction-Harvest-Robot",
        "模型权重转换仓库：https://github.com/bertholetnadine159-prog/Model-weight-conversion",
        "hobot_websocket 官方实现：https://github.com/D-Robotics/hobot_websocket",
    ]),
]


# ============================================================================
# DOCX 生成
# ============================================================================
def _set_run_font(run, size: float, bold: bool = False, color: str | None = None,
                  ascii_font: str = "Calibri", east_font: str = "微软雅黑",
                  mono: bool = False) -> None:
    from docx.shared import Pt, RGBColor

    run.font.name = "Consolas" if mono else ascii_font
    run.font.size = Pt(size)
    run.font.bold = bold
    if color:
        run.font.color.rgb = RGBColor.from_string(color.lstrip("#"))
    r_pr = run._element.get_or_add_rPr()
    r_fonts = r_pr.rFonts
    if r_fonts is None:
        r_fonts = r_pr.get_or_add_rFonts()
    r_fonts.set(
        "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}eastAsia",
        east_font,
    )


def _shade_paragraph(paragraph, fill: str) -> None:
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    p_pr = paragraph._p.get_or_add_pPr()
    shading = OxmlElement("w:shd")
    shading.set(qn("w:val"), "clear")
    shading.set(qn("w:color"), "auto")
    shading.set(qn("w:fill"), fill.lstrip("#"))
    p_pr.append(shading)


def _left_border(paragraph, color: str = "2E74B5", size: int = 18) -> None:
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    p_pr = paragraph._p.get_or_add_pPr()
    borders = OxmlElement("w:pBdr")
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), str(size))
    left.set(qn("w:space"), "4")
    left.set(qn("w:color"), color)
    borders.append(left)
    p_pr.append(borders)


def _setup_numbering(document) -> None:
    """注入两种真实编号定义：0=圆点列表，1=十进制列表。"""
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    def make_level(fmt_value: str, text_value: str) -> Any:
        lvl = OxmlElement("w:lvl")
        lvl.set(qn("w:ilvl"), "0")
        start = OxmlElement("w:start")
        start.set(qn("w:val"), "1")
        fmt = OxmlElement("w:numFmt")
        fmt.set(qn("w:val"), fmt_value)
        text = OxmlElement("w:lvlText")
        text.set(qn("w:val"), text_value)
        p_pr = OxmlElement("w:pPr")
        ind = OxmlElement("w:ind")
        ind.set(qn("w:left"), "720")
        ind.set(qn("w:hanging"), "360")
        p_pr.append(ind)
        lvl.extend([start, fmt, text, p_pr])
        return lvl

    numbering = document.part.numbering_part.element
    abstract_bullet = OxmlElement("w:abstractNum")
    abstract_bullet.set(qn("w:abstractNumId"), "0")
    abstract_bullet.append(make_level("bullet", "•"))

    abstract_decimal = OxmlElement("w:abstractNum")
    abstract_decimal.set(qn("w:abstractNumId"), "1")
    abstract_decimal.append(make_level("decimal", "%1."))

    num_bullet = OxmlElement("w:num")
    num_bullet.set(qn("w:numId"), "100")
    bullet_ref = OxmlElement("w:abstractNumId")
    bullet_ref.set(qn("w:val"), "0")
    num_bullet.append(bullet_ref)
    num_decimal = OxmlElement("w:num")
    num_decimal.set(qn("w:numId"), "101")
    decimal_ref = OxmlElement("w:abstractNumId")
    decimal_ref.set(qn("w:val"), "1")
    num_decimal.append(decimal_ref)
    numbering.extend([abstract_bullet, abstract_decimal, num_bullet, num_decimal])


def _add_list_item(document, text: str, num_id: str = "100") -> None:
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_after = 8
    paragraph.paragraph_format.line_spacing = 1.167
    p_pr = paragraph._p.get_or_add_pPr()
    num_pr = OxmlElement("w:numPr")
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    num_id_el = OxmlElement("w:numId")
    num_id_el.set(qn("w:val"), num_id)
    num_pr.extend([ilvl, num_id_el])
    p_pr.append(num_pr)
    _set_run_font(paragraph.add_run(text), 11)


def _set_cell_background(cell, fill: str) -> None:
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    tc_pr = cell._tc.get_or_add_tcPr()
    shading = OxmlElement("w:shd")
    shading.set(qn("w:val"), "clear")
    shading.set(qn("w:color"), "auto")
    shading.set(qn("w:fill"), fill.lstrip("#"))
    tc_pr.append(shading)


def _set_table_geometry(table, widths_dxa: list[int]) -> None:
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    tbl = table._tbl
    tbl_pr = tbl.tblPr
    layout = OxmlElement("w:tblLayout")
    layout.set(qn("w:type"), "fixed")
    tbl_pr.append(layout)

    grid = tbl.find(qn("w:tblGrid"))
    if grid is not None:
        tbl.remove(grid)
    grid = OxmlElement("w:tblGrid")
    for width in widths_dxa:
        grid_col = OxmlElement("w:gridCol")
        grid_col.set(qn("w:w"), str(width))
        grid.append(grid_col)
    tbl.insert(list(tbl).index(tbl_pr) + 1, grid)

    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths_dxa)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")

    for row in table.rows:
        for cell, width in zip(row.cells, widths_dxa):
            from docx.shared import Emu
            cell.width = Emu(width * 635)
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")
            margins = OxmlElement("w:tcMar")
            for side, value in (("top", "80"), ("bottom", "80"), ("start", "120"), ("end", "120")):
                el = OxmlElement(f"w:{side}")
                el.set(qn("w:w"), value)
                el.set(qn("w:type"), "dxa")
                margins.append(el)
            tc_pr.append(margins)


def build_docx(path: Path) -> None:
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
    from docx.shared import Inches, Pt

    document = Document()
    section = document.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal = document.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal_r_pr = normal.element.get_or_add_rPr()
    normal_r_fonts = normal_r_pr.get_or_add_rFonts()
    normal_r_fonts.set(
        "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}eastAsia",
        "微软雅黑",
    )
    normal.paragraph_format.space_after = 6
    normal.paragraph_format.line_spacing = 1.10

    _setup_numbering(document)

    for kind, payload in CONTENT:
        if kind == "title":
            paragraph = document.add_paragraph()
            paragraph.paragraph_format.space_before = 0
            paragraph.paragraph_format.space_after = 4
            _set_run_font(paragraph.add_run(payload), 20, bold=True, color=INK_BLUE)
        elif kind == "subtitle":
            paragraph = document.add_paragraph()
            paragraph.paragraph_format.space_after = 4
            _set_run_font(paragraph.add_run(payload), 12, color=GRAY)
        elif kind == "meta":
            paragraph = document.add_paragraph()
            paragraph.paragraph_format.space_after = 18
            _set_run_font(paragraph.add_run(payload), 9.5, color=GRAY)
        elif kind == "h1":
            paragraph = document.add_paragraph(style="Heading 1")
            paragraph.paragraph_format.space_before = 16
            paragraph.paragraph_format.space_after = 8
            for run in paragraph.runs:
                run.text = ""
            _set_run_font(paragraph.add_run(payload), 16, bold=True, color=HEADING_BLUE)
        elif kind == "h2":
            paragraph = document.add_paragraph(style="Heading 2")
            paragraph.paragraph_format.space_before = 12
            paragraph.paragraph_format.space_after = 6
            for run in paragraph.runs:
                run.text = ""
            _set_run_font(paragraph.add_run(payload), 13, bold=True, color=HEADING_BLUE)
        elif kind == "para":
            paragraph = document.add_paragraph()
            _set_run_font(paragraph.add_run(payload), 11)
        elif kind == "bullets":
            for item in payload:
                _add_list_item(document, item, num_id="100")
        elif kind == "code":
            paragraph = document.add_paragraph()
            paragraph.paragraph_format.space_after = 8
            paragraph.paragraph_format.left_indent = 180
            paragraph.paragraph_format.line_spacing = 1.0
            _shade_paragraph(paragraph, "F4F6F9")
            _left_border(paragraph)
            for index, line in enumerate(payload.splitlines()):
                if index:
                    paragraph.add_run().add_break()
                _set_run_font(paragraph.add_run(line), 9, mono=True)
        elif kind == "note":
            paragraph = document.add_paragraph()
            paragraph.paragraph_format.space_before = 4
            paragraph.paragraph_format.space_after = 8
            paragraph.paragraph_format.left_indent = 180
            _shade_paragraph(paragraph, "FDF3E7")
            _left_border(paragraph, color="C77A1B")
            _set_run_font(paragraph.add_run(payload), 10)
        elif kind == "table":
            headers = payload["headers"]
            widths = payload["widths"]
            rows = payload["rows"]
            table = document.add_table(rows=1 + len(rows), cols=len(headers))
            table.style = "Table Grid"
            _set_table_geometry(table, widths)
            for col, header in enumerate(headers):
                cell = table.rows[0].cells[col]
                _set_cell_background(cell, HEADER_FILL)
                cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                paragraph = cell.paragraphs[0]
                paragraph.paragraph_format.space_after = 0
                _set_run_font(paragraph.add_run(header), 10, bold=True, color=INK_BLUE)
            for row_index, row in enumerate(rows):
                for col, value in enumerate(row):
                    cell = table.rows[row_index + 1].cells[col]
                    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                    paragraph = cell.paragraphs[0]
                    paragraph.paragraph_format.space_after = 0
                    _set_run_font(paragraph.add_run(value), 10)
            document.add_paragraph().paragraph_format.space_after = 2

    for section in document.sections:
        footer = section.footer
        footer_paragraph = footer.paragraphs[0]
        footer_paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        _set_run_font(footer_paragraph.add_run("SeaUI 软件框架说明 · 第 "), 8.5, color=GRAY)
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn
        run = footer_paragraph.add_run()
        fld_char1 = OxmlElement("w:fldChar")
        fld_char1.set(qn("w:fldCharType"), "begin")
        instr = OxmlElement("w:instrText")
        instr.set(qn("xml:space"), "preserve")
        instr.text = "PAGE"
        fld_char2 = OxmlElement("w:fldChar")
        fld_char2.set(qn("w:fldCharType"), "end")
        run._r.append(fld_char1)
        run._r.append(instr)
        run._r.append(fld_char2)
        _set_run_font(run, 8.5, color=GRAY)
        _set_run_font(footer_paragraph.add_run(" 页"), 8.5, color=GRAY)

    document.save(path)


# ============================================================================
# PDF 生成（reportlab，与 DOCX 同内容同视觉体系）
# ============================================================================
def build_pdf(path: Path) -> None:
    from reportlab.lib import colors
    from reportlab.lib.enums import TA_LEFT, TA_RIGHT
    from reportlab.lib.pagesizes import letter
    from reportlab.lib.styles import ParagraphStyle
    from reportlab.lib.units import inch
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.ttfonts import TTFont
    from reportlab.platypus import (
        KeepTogether,
        PageBreak,
        Paragraph,
        SimpleDocTemplate,
        Spacer,
        Table,
        TableStyle,
    )

    pdfmetrics.registerFont(TTFont("YaHei", r"C:\Windows\Fonts\msyh.ttc", subfontIndex=0))
    pdfmetrics.registerFont(TTFont("YaHeiBold", r"C:\Windows\Fonts\msyhbd.ttc", subfontIndex=0))
    pdfmetrics.registerFont(TTFont("Consolas", r"C:\Windows\Fonts\consola.ttf"))
    pdfmetrics.registerFontFamily("YaHei", normal="YaHei", bold="YaHeiBold", italic="YaHei", boldItalic="YaHeiBold")

    blue = colors.HexColor(HEADING_BLUE)
    dark_blue = colors.HexColor(INK_BLUE)
    gray = colors.HexColor(GRAY)

    styles = {
        "title": ParagraphStyle("title", fontName="YaHeiBold", fontSize=20, leading=25, textColor=dark_blue, spaceAfter=4),
        "subtitle": ParagraphStyle("subtitle", fontName="YaHei", fontSize=12, leading=16, textColor=gray, spaceAfter=4),
        "meta": ParagraphStyle("meta", fontName="YaHei", fontSize=9.5, leading=13, textColor=gray, spaceAfter=14),
        "h1": ParagraphStyle("h1", fontName="YaHeiBold", fontSize=16, leading=20, textColor=blue, spaceBefore=14, spaceAfter=7),
        "h2": ParagraphStyle("h2", fontName="YaHeiBold", fontSize=13, leading=17, textColor=blue, spaceBefore=10, spaceAfter=5),
        "body": ParagraphStyle("body", fontName="YaHei", fontSize=11, leading=13.2, spaceAfter=6, alignment=TA_LEFT),
        "bullet": ParagraphStyle("bullet", fontName="YaHei", fontSize=11, leading=13.2, spaceAfter=5, leftIndent=18, bulletIndent=6),
        "code": ParagraphStyle("code", fontName="Consolas", fontSize=8.5, leading=11.5, spaceAfter=8, leftIndent=10, rightIndent=10),
        "note": ParagraphStyle("note", fontName="YaHei", fontSize=10, leading=13.5, spaceAfter=8, leftIndent=10, rightIndent=10, textColor=colors.HexColor("#7A4A00")),
        "cell": ParagraphStyle("cell", fontName="YaHei", fontSize=10, leading=12.5),
        "cellhead": ParagraphStyle("cellhead", fontName="YaHeiBold", fontSize=10, leading=12.5, textColor=dark_blue),
        "footer": ParagraphStyle("footer", fontName="YaHei", fontSize=8.5, leading=10, textColor=gray, alignment=TA_RIGHT),
    }

    story = []
    total_width = letter[0] - 2 * inch

    for kind, payload in CONTENT:
        if kind == "title":
            story.append(Paragraph(payload, styles["title"]))
        elif kind == "subtitle":
            story.append(Paragraph(payload, styles["subtitle"]))
        elif kind == "meta":
            story.append(Paragraph(payload, styles["meta"]))
        elif kind == "h1":
            story.append(Paragraph(payload, styles["h1"]))
        elif kind == "h2":
            story.append(Paragraph(payload, styles["h2"]))
        elif kind == "para":
            story.append(Paragraph(payload, styles["body"]))
        elif kind == "bullets":
            for item in payload:
                story.append(Paragraph(item, styles["bullet"], bulletText="•"))
        elif kind == "code":
            lines = payload.splitlines()
            story.append(Table(
                [[Paragraph("<br/>".join(lines), styles["code"])]],
                colWidths=[total_width - 20],
                style=TableStyle([
                    ("BACKGROUND", (0, 0), (-1, -1), colors.HexColor("#F4F6F9")),
                    ("LINEBEFORE", (0, 0), (0, -1), 1.5, blue),
                    ("LEFTPADDING", (0, 0), (-1, -1), 10),
                    ("RIGHTPADDING", (0, 0), (-1, -1), 10),
                    ("TOPPADDING", (0, 0), (-1, -1), 8),
                    ("BOTTOMPADDING", (0, 0), (-1, -1), 8),
                ]),
            ))
        elif kind == "note":
            story.append(Table(
                [[Paragraph(payload, styles["note"])]],
                colWidths=[total_width - 20],
                style=TableStyle([
                    ("BACKGROUND", (0, 0), (-1, -1), colors.HexColor("#FDF3E7")),
                    ("LINEBEFORE", (0, 0), (0, -1), 1.5, colors.HexColor("#C77A1B")),
                    ("LEFTPADDING", (0, 0), (-1, -1), 10),
                    ("RIGHTPADDING", (0, 0), (-1, -1), 10),
                    ("TOPPADDING", (0, 0), (-1, -1), 7),
                    ("BOTTOMPADDING", (0, 0), (-1, -1), 7),
                ]),
            ))
        elif kind == "table":
            headers = payload["headers"]
            rows = payload["rows"]
            scale = total_width / sum(payload["widths"])
            widths = [w * scale for w in payload["widths"]]
            data = [[Paragraph(h, styles["cellhead"]) for h in headers]]
            for row in rows:
                data.append([Paragraph(cell, styles["cell"]) for cell in row])
            table = Table(data, colWidths=widths, repeatRows=1, hAlign="LEFT")
            table.setStyle(TableStyle([
                ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor(HEADER_FILL)),
                ("GRID", (0, 0), (-1, -1), 0.5, colors.HexColor(BORDER_GRAY)),
                ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
                ("LEFTPADDING", (0, 0), (-1, -1), 7),
                ("RIGHTPADDING", (0, 0), (-1, -1), 7),
                ("TOPPADDING", (0, 0), (-1, -1), 6),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 6),
            ]))
            story.append(table)
            story.append(Spacer(1, 8))

    def footer(canvas, doc):
        canvas.saveState()
        canvas.setFont("YaHei", 8.5)
        canvas.setFillColor(gray)
        canvas.drawRightString(letter[0] - inch, 0.55 * inch, f"SeaUI 软件框架说明 · 第 {doc.page} 页")
        canvas.restoreState()

    doc = SimpleDocTemplate(
        str(path),
        pagesize=letter,
        leftMargin=inch,
        rightMargin=inch,
        topMargin=inch,
        bottomMargin=inch,
        title="SeaUI 软件框架说明",
        author="SeaUI Team",
    )
    doc.build(story, onFirstPage=footer, onLaterPages=footer)


if __name__ == "__main__":
    build_docx(DOCX_PATH)
    build_pdf(PDF_PATH)
    print("built:", DOCX_PATH)
    print("built:", PDF_PATH)
